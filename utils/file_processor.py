import os
import io
import streamlit as st
from typing import List, Dict, Any, Optional
import PyPDF2
from docx import Document
import markdown
from utils.logger import get_logger
from utils.file_metadata import FileMetadataManager

logger = get_logger(__name__)

class FileProcessor:
    """Class để xử lý các loại file khác nhau"""

    SUPPORTED_EXTENSIONS = {
        '.txt': 'text/plain',
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.md': 'text/markdown'
    }

    @staticmethod
    def format_file_size(size_bytes: int) -> str:
        """Format file size with appropriate unit (Bytes, KB, MB)"""
        if size_bytes < 1024:
            return f"{size_bytes} Bytes"
        elif size_bytes < 1024 * 1024:
            size_kb = size_bytes / 1024
            return f"{size_kb:.1f}KB"
        else:
            size_mb = size_bytes / (1024 * 1024)
            return f"{size_mb:.1f}MB"
    
    @staticmethod
    def extract_text_from_file(file_content: bytes, filename: str) -> str:
        """Trích xuất text từ file content"""
        file_ext = os.path.splitext(filename)[1].lower()
        logger.debug(f"Extracting text from {filename} (type: {file_ext})")

        try:
            if file_ext == '.txt':
                text = file_content.decode('utf-8')
                logger.debug(f"Extracted {len(text)} chars from TXT file: {filename}")
                return text
            
            elif file_ext == '.pdf':
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                logger.debug(f"Extracted {len(text)} chars from PDF file ({len(pdf_reader.pages)} pages): {filename}")
                return text
            
            elif file_ext == '.docx':
                doc = Document(io.BytesIO(file_content))
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                logger.debug(f"Extracted {len(text)} chars from DOCX file ({len(doc.paragraphs)} paragraphs): {filename}")
                return text
            
            elif file_ext == '.md':
                # Convert markdown to plain text
                md_text = file_content.decode('utf-8')
                html = markdown.markdown(md_text)
                # Simple HTML to text conversion
                import re
                text = re.sub(r'<[^>]+>', '', html)
                logger.debug(f"Extracted {len(text)} chars from MD file: {filename}")
                return text

            else:
                logger.error(f"Unsupported file type: {file_ext} for file: {filename}")
                raise ValueError(f"Unsupported file type: {file_ext}")

        except Exception as e:
            st.error(f"Lỗi khi xử lý file {filename}: {str(e)}")
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            return ""
    
    @staticmethod
    def validate_file(file) -> Dict[str, Any]:
        """Validate file upload"""
        result = {
            'valid': True,
            'error': None,
            'size_mb': 0
        }
        
        # Check file size
        file_size = len(file.getvalue())
        result['size_mb'] = file_size / (1024 * 1024)
        
        if result['size_mb'] > 200:  # 200MB limit
            result['valid'] = False
            result['error'] = f"File quá lớn ({result['size_mb']:.1f}MB). Giới hạn 200MB."
            return result
        
        # Check file extension
        file_ext = os.path.splitext(file.name)[1].lower()
        if file_ext not in FileProcessor.SUPPORTED_EXTENSIONS:
            result['valid'] = False
            result['error'] = f"Loại file không được hỗ trợ: {file_ext}. Chỉ hỗ trợ: {', '.join(FileProcessor.SUPPORTED_EXTENSIONS.keys())}"
        
        return result
    
    @staticmethod
    def process_uploaded_files(uploaded_files: List, save_to_disk: bool = True, uploads_dir: str = "./uploads",
                               metadata_file: str = ".metadata.json", hash_algorithm: str = "sha256",
                               project_id: str = None) -> Dict[str, Any]:
        """
        Xử lý danh sách file đã upload với duplicate detection
        
        Args:
            uploaded_files: List of uploaded file objects
            save_to_disk: Whether to save files to disk
            uploads_dir: Base uploads directory (will be modified if project_id is provided)
            metadata_file: Name of metadata file
            hash_algorithm: Hash algorithm to use
            project_id: Project identifier for project-scoped storage
            
        Returns: {
            'processed_files': List[Dict],
            'stats': {
                'new': int, 'duplicates': int, 'updated': int, 'errors': int
            },
            'duplicates': List[Dict]  # Thông tin các file duplicate
        }
        """
        from config import Config
        
        # Use project-scoped directory if project_id is provided
        if project_id:
            uploads_dir = Config.get_project_uploads_dir(project_id)
            logger.info(f"Using project-scoped uploads directory: {uploads_dir}")
        
        processed_files = []
        duplicate_files = []
        stats = {'new': 0, 'duplicates': 0, 'updated': 0, 'errors': 0}

        logger.info(f"Processing {len(uploaded_files)} uploaded files with duplicate detection")

        # Create uploads directory if needed
        if save_to_disk and not os.path.exists(uploads_dir):
            os.makedirs(uploads_dir)
            logger.info(f"Created uploads directory: {uploads_dir}")

        # Initialize metadata manager
        metadata_manager = FileMetadataManager(uploads_dir, metadata_file)

        for uploaded_file in uploaded_files:
            # Validate file
            validation = FileProcessor.validate_file(uploaded_file)
            if not validation['valid']:
                st.error(f"File {uploaded_file.name}: {validation['error']}")
                logger.warning(f"File validation failed for {uploaded_file.name}: {validation['error']}")
                stats['errors'] += 1
                continue

            # Get file content and compute hash
            file_content = uploaded_file.getvalue()
            file_hash = FileMetadataManager.compute_file_hash(file_content, hash_algorithm)
            file_size_bytes = len(file_content)

            # Check for duplicates
            duplicate_check = metadata_manager.check_duplicate(
                uploaded_file.name,
                file_hash,
                file_size_bytes
            )

            # Handle exact duplicates (skip)
            if duplicate_check['is_duplicate'] and duplicate_check['duplicate_type'] == 'exact':
                duplicate_files.append({
                    'name': uploaded_file.name,
                    'type': duplicate_check['duplicate_type'],
                    'message': duplicate_check['message'],
                    'existing_file': duplicate_check['existing_file'],
                    'hash': file_hash[:16] + '...',
                    'size_formatted': FileProcessor.format_file_size(file_size_bytes)
                })
                stats['duplicates'] += 1
                logger.info(f"Skipped exact duplicate: {uploaded_file.name}")
                continue

            # Handle content duplicates (different filename, same content) - skip
            if duplicate_check['is_duplicate'] and duplicate_check['duplicate_type'] == 'content':
                duplicate_files.append({
                    'name': uploaded_file.name,
                    'type': duplicate_check['duplicate_type'],
                    'message': duplicate_check['message'],
                    'existing_file': duplicate_check['existing_file'],
                    'hash': file_hash[:16] + '...',
                    'size_formatted': FileProcessor.format_file_size(file_size_bytes)
                })
                stats['duplicates'] += 1
                logger.info(f"Skipped content duplicate: {uploaded_file.name} (matches {duplicate_check['existing_file']})")
                continue

            # Extract text
            text = FileProcessor.extract_text_from_file(file_content, uploaded_file.name)

            if text.strip():
                # Save file to disk
                if save_to_disk:
                    file_path = os.path.join(uploads_dir, uploaded_file.name)
                    with open(file_path, 'wb') as f:
                        f.write(file_content)
                    logger.info(f"Saved file to disk: {file_path}")

                # Add to metadata
                metadata_manager.add_file(uploaded_file.name, file_hash, file_size_bytes, processed=True)

                # Track stats
                if duplicate_check['duplicate_type'] == 'updated':
                    stats['updated'] += 1
                    status = 'updated'
                else:
                    stats['new'] += 1
                    status = 'new'

                processed_files.append({
                    'name': uploaded_file.name,
                    'content': text,
                    'size_mb': validation['size_mb'],
                    'size_bytes': file_size_bytes,
                    'size_formatted': FileProcessor.format_file_size(file_size_bytes),
                    'type': os.path.splitext(uploaded_file.name)[1].lower(),
                    'hash': file_hash[:16] + '...',  # Short hash for display
                    'hash_full': file_hash,
                    'status': status
                })
                logger.info(f"Successfully processed file ({status}): {uploaded_file.name} ({FileProcessor.format_file_size(file_size_bytes)})")
            else:
                st.warning(f"Không thể trích xuất text từ file: {uploaded_file.name}")
                logger.warning(f"No text extracted from file: {uploaded_file.name}")
                stats['errors'] += 1

        logger.info(f"Processing complete - New: {stats['new']}, Updated: {stats['updated']}, Duplicates: {stats['duplicates']}, Errors: {stats['errors']}")

        return {
            'processed_files': processed_files,
            'stats': stats,
            'duplicates': duplicate_files
        }
    
    @staticmethod
    def get_file_preview(content: str, max_chars: int = 500) -> str:
        """Lấy preview của file content"""
        if len(content) <= max_chars:
            return content
        return content[:max_chars] + "..."

    @staticmethod
    def check_uploads_directory(uploads_dir: str = "./uploads", project_id: str = None) -> Dict[str, Any]:
        """
        Check if uploads directory exists and has files
        
        Args:
            uploads_dir: Base uploads directory
            project_id: Project identifier for project-scoped storage
            
        Returns:
            Dictionary with directory status and file list
        """
        from config import Config
        
        # Use project-scoped directory if project_id is provided
        if project_id:
            uploads_dir = Config.get_project_uploads_dir(project_id)
        
        result = {
            'exists': False,
            'has_files': False,
            'file_count': 0,
            'files': []
        }

        if os.path.exists(uploads_dir):
            result['exists'] = True
            files = [f for f in os.listdir(uploads_dir) if os.path.isfile(os.path.join(uploads_dir, f))]
            result['file_count'] = len(files)
            result['has_files'] = len(files) > 0
            result['files'] = files
            logger.debug(f"Uploads directory check: {result['file_count']} files found")
        else:
            logger.debug(f"Uploads directory does not exist: {uploads_dir}")

        return result
